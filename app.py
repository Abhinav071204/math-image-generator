import streamlit as st
import os, re, json, io, zipfile
from pathlib import Path
import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches


# ============================================================
# PAGE CONFIG
# ============================================================
st.set_page_config(
    page_title='Math Image Generator',
    page_icon='📐',
    layout='wide',
    initial_sidebar_state='expanded'
)

st.markdown("""
<style>
.stButton>button { border-radius: 8px; }
.stDownloadButton>button { border-radius: 8px; background: #2e7d32; color: white; }
</style>
""", unsafe_allow_html=True)

# ============================================================
# CONFIG
# ============================================================
STYLE_RULES = {
    'graphs': {
        'axis_pad_units'   : 0.5,
        'line_width_pt'    : 1.5,
        'gridline_width_pt': 0.5,
        'gridline_color'   : '#cccccc',
    },
}
GRAYSCALE_PATTERNS = ['-', '--', ':', '-.']
COLOR_PALETTE      = ['blue', 'darkorange', 'green', 'red', 'purple']

# ============================================================
# OPTIONAL AI PARSING (off by default — costs money only if a key is supplied)
# ============================================================
AI_SYSTEM_PROMPT = """You convert a math-graph image prompt into STRICT JSON. Output ONLY valid JSON, no markdown fences, no commentary.

Detect exactly one of two graph types and return ONLY the matching schema.

TYPE "line_graph":
{
  "type": "line_graph",
  "title": string or null,
  "x_label": string,
  "y_label": string,
  "x_min": number, "x_max": number,
  "y_min": number, "y_max": number,
  "grade_band": "3-5" | "6-8" | "9-12" | null,
  "lines": [
    {"equation": string, "label": string or null, "point1": [x,y], "point2": [x,y], "style": "solid"|"dashed"|"dotted"}
  ],
  "intersection": [x, y] or null
}

TYPE "scatter_grid":
{
  "type": "scatter_grid",
  "x_label": string, "y_label": string,
  "x_min": number, "x_max": number, "y_min": number, "y_max": number,
  "grade_band": "3-5" | "6-8" | "9-12" | null,
  "plots": { "A": [[x,y],[x,y],...], "B": [...] }
}

Rules:
- If axis range isn't stated, default x:-2 to 8, y:-2 to 8 for line_graph; 0 to 10 for scatter_grid.
- grade_band: "grade 3-5"/"elementary" -> "3-5"; "grade 6-8"/"middle school" -> "6-8"; "grade 9-12"/"high school" -> "9-12". Otherwise null.
- Do your best with incomplete prompts rather than failing."""


def ai_parse_prompt(prompt, api_key):
    """Returns (data_dict, error). Never raises — failures return (None, msg)
    so callers can fall back to the free regex parser cleanly."""
    try:
        import anthropic
    except ImportError:
        return None, "anthropic package not installed"
    if not api_key:
        return None, "No API key provided"
    try:
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=1500,
            system=AI_SYSTEM_PROMPT,
            messages=[{"role": "user", "content": prompt}]
        )
        text = resp.content[0].text.strip()
        text = re.sub(r'^```(?:json)?\s*|\s*```$', '', text.strip())
        return json.loads(text), None
    except json.JSONDecodeError as e:
        return None, f"Model did not return valid JSON: {e}"
    except Exception as e:
        return None, f"AI call failed: {e}"


def normalize_ai_line_params(data):
    grade_band = data.get('grade_band')
    color_mode = resolve_color_mode(grade_band)
    linestyle_map = {'solid': '-', 'dashed': '--', 'dotted': ':', 'dash-dot': '-.'}
    lines_out = []
    for line in data.get('lines', []):
        p1, p2 = line.get('point1'), line.get('point2')
        if not p1 or not p2:
            continue
        x1, y1 = p1
        x2, y2 = p2
        if x2 == x1:
            slope, intercept = None, None
        else:
            slope = (y2 - y1) / (x2 - x1)
            intercept = y1 - slope * x1
        lines_out.append({
            'equation'      : line.get('label') or line.get('equation') or 'line',
            'slope'         : slope,
            'intercept'     : intercept,
            'points'        : [(x1, y1), (x2, y2)],
            'linestyle'     : linestyle_map.get(line.get('style', 'solid'), '-'),
            'style_explicit': True,
        })
    ds = get_dataset_styles(len(lines_out), color_mode)
    for idx, line in enumerate(lines_out):
        line['color'] = ds[idx][0]
    intersection = data.get('intersection')
    return {
        'title'       : data.get('title') or '',
        'x_label'     : data.get('x_label', 'x'),
        'y_label'     : data.get('y_label', 'y'),
        'x_min'       : data.get('x_min', -2),
        'x_max'       : data.get('x_max', 8),
        'y_min'       : data.get('y_min', -2),
        'y_max'       : data.get('y_max', 8),
        'lines'       : lines_out,
        'intersection': tuple(intersection) if intersection else None,
        'grade_band'  : grade_band,
        'color_mode'  : color_mode,
    }


def normalize_ai_scatter_params(data):
    grade_band = data.get('grade_band')
    plots = {k: [tuple(pt) for pt in v] for k, v in data.get('plots', {}).items()}
    return {
        'x_label'   : data.get('x_label', 'x'),
        'y_label'   : data.get('y_label', 'y'),
        'x_min'     : data.get('x_min', 0),
        'x_max'     : data.get('x_max', 10),
        'y_min'     : data.get('y_min', 0),
        'y_max'     : data.get('y_max', 10),
        'plots'     : plots,
        'grade_band': grade_band,
        'color_mode': resolve_color_mode(grade_band),
        'plot_descriptions': {},
    }


# ============================================================
# REGEX PARSERS (always free, default path)
# ============================================================
_LOWER_WORDS = {'a','an','the','in','on','at','to','for','of','with','by','and','but','or','nor','so','yet'}

def to_title_case(text):
    words = text.split(' ')
    out = []
    for i, w in enumerate(words):
        lw = w.lower()
        out.append(lw if i != 0 and lw in _LOWER_WORDS else lw[:1].upper() + lw[1:])
    return ' '.join(out)

def to_sentence_case(text):
    if not text: return text
    return text.strip()[:1].upper() + text.strip()[1:]

def extract_grade_band(prompt):
    p = prompt.lower()
    if re.search(r'grade\s*(3|4|5)|grades?\s*3\s*-\s*5|elementary', p): return '3-5'
    if re.search(r'grade\s*(6|7|8)|grades?\s*6\s*-\s*8|middle', p):     return '6-8'
    if re.search(r'grade\s*(9|10|11|12)|grades?\s*9\s*-\s*12|high\s*school', p): return '9-12'
    return None

def resolve_color_mode(grade_band):
    return 'color' if grade_band == '3-5' else 'grayscale'

def get_dataset_styles(n, color_mode):
    styles = []
    for i in range(n):
        if color_mode == 'color':
            styles.append((COLOR_PALETTE[i % len(COLOR_PALETTE)], '-'))
        else:
            styles.append(('black', GRAYSCALE_PATTERNS[i % len(GRAYSCALE_PATTERNS)]))
    return styles

def resolve_font(grade_band):
    available = {f.name for f in fm.fontManager.ttflist}
    wanted = 'Poppins' if grade_band == '3-5' else 'DejaVu Sans'
    return wanted if wanted in available else 'DejaVu Sans'

def detect_prompt_type(prompt):
    p = prompt.lower()
    if re.search(r'scatter|2x2|four\s+(scatter|plot)', p):
        return 'scatter_grid'
    if re.search(r'\bplot\s+[a-z0-9]\s*[:\-]|\bplot\s+[a-z0-9]\s+points?', p, re.I):
        return 'scatter_grid'
    if re.search(r'coordinate\s+plane|passes?\s+through|passing\s+through|goes?\s+through|intersect|y\s*=|slope|line', p):
        return 'line_graph'
    return 'unknown'

def parse_line_prompt(prompt):
    p = prompt
    for ch in ['−', '–', '—', '\u2212', '\u2013', '\u2014']:
        p = p.replace(ch, '-')
    p_flat = p.replace('\n', ' ').strip()
    grade_band = extract_grade_band(p_flat)
    params = {
        'title': '', 'x_label': 'x', 'y_label': 'y',
        'x_min': -2, 'x_max': 8, 'y_min': -2, 'y_max': 8,
        'lines': [], 'intersection': None,
        'grade_band': grade_band,
        'color_mode': resolve_color_mode(grade_band),
    }
    t = re.search(r'titled?\s*:?\s*"([^"]+)"', p_flat, re.I)
    if t: params['title'] = to_title_case(t.group(1))
    xl = re.search(r'x-axis\s*:?\s*(?:labeled?|as)?\s*"([^"]+)"', p_flat, re.I)
    yl = re.search(r'y-axis\s*:?\s*(?:labeled?|as)?\s*"([^"]+)"', p_flat, re.I)
    if xl: params['x_label'] = to_sentence_case(xl.group(1))
    if yl: params['y_label'] = to_sentence_case(yl.group(1))
    combined = re.search(r'(?:x-?\s*and\s*y-?axes?|axes?)\s+(?:from|range)\s+(-?\d+(?:\.\d+)?)\s+to\s+(-?\d+(?:\.\d+)?)', p_flat, re.I)
    if combined:
        lo, hi = float(combined.group(1)), float(combined.group(2))
        params['x_min'] = params['y_min'] = lo
        params['x_max'] = params['y_max'] = hi
    xr = re.search(r'x-axis[^.\n]*?(?:from|range)\s+(-?\d+(?:\.\d+)?)\s+to\s+(-?\d+(?:\.\d+)?)', p_flat, re.I)
    yr = re.search(r'y-axis[^.\n]*?(?:from|range)\s+(-?\d+(?:\.\d+)?)\s+to\s+(-?\d+(?:\.\d+)?)', p_flat, re.I)
    if xr: params['x_min'], params['x_max'] = float(xr.group(1)), float(xr.group(2))
    if yr: params['y_min'], params['y_max'] = float(yr.group(1)), float(yr.group(2))
    for m in re.finditer(r'x-axis.*?\(extend to\s+(-?\d+(?:\.\d+)?)\)', p_flat, re.I):
        params['x_max'] = float(m.group(1))
    for m in re.finditer(r'y-axis.*?\(extend to\s+(-?\d+(?:\.\d+)?)\)', p_flat, re.I):
        params['y_max'] = float(m.group(1))

    linestyles_map = {'solid': '-', 'dashed': '--', 'dotted': ':', 'dash-dot': '-.'}

    def pts_to_slope(x1, y1, x2, y2):
        if x2 == x1: return None, None
        s = (y2 - y1) / (x2 - x1)
        return s, y1 - s * x1

    def parse_equation(text):
        text = text.replace(' ', '').lower()
        m = re.match(r'y=(-?\d*\.?\d*)x([+-]\d+\.?\d*)?$', text)
        if m:
            coef = m.group(1)
            if coef in ('', '+'):   slope = 1.0
            elif coef == '-':       slope = -1.0
            else:                   slope = float(coef)
            intercept = float(m.group(2)) if m.group(2) else 0.0
            return slope, intercept
        m = re.match(r'y=(-?\d+\.?\d*)$', text)
        if m:
            return 0.0, float(m.group(1))
        return None

    block_spans = [m.start() for m in re.finditer(
        r'Dataset\s+\d+[^\n]*?:|Line\s+[A-Z]\s*:|(?:the\s+)?line\s+y\s*=', p, re.I)]
    blocks = []
    if block_spans:
        for i, s in enumerate(block_spans):
            e = block_spans[i + 1] if i + 1 < len(block_spans) else len(p)
            blocks.append(p[s:e])
    else:
        blocks = [p]

    def extract_name(block):
        m = re.match(r'Dataset\s+\d+\s*[—\-:]\s*([^\n(]+?)(?:\s*\(|:|\n)', block, re.I)
        if m: return m.group(1).strip(' —-:')
        m = re.match(r'(Line\s+[A-Z])', block, re.I)
        if m: return m.group(1).strip()
        m = re.search(r'\by\s*=\s*-?\s*\d*\.?\d*\s*x\s*(?:[+\-]\s*\d+\.?\d*)?', block, re.I)
        if m:
            eq = re.sub(r'\s+', ' ', m.group(0))
            return re.sub(r'\s*([=+\-])\s*', r' \1 ', eq).strip()
        return None

    def extract_style(block):
        m = re.search(r'\((solid|dashed|dotted|dash-dot)\s*(?:line)?\)', block, re.I)
        return m.group(1).lower() if m else None

    def extract_coord_pairs(block):
        pairs = []
        for x, y in re.findall(r'\(\s*(-?\d+(?:\.\d+)?)\s*,\s*(-?\d+(?:\.\d+)?)\s*\)', block):
            pt = (float(x), float(y))
            if not pairs or pairs[-1] != pt:
                pairs.append(pt)
        if pairs:
            return pairs
        for x, y in re.findall(r'x\s*=\s*(-?\d+(?:\.\d+)?)\D+?y\s*=\s*(-?\d+(?:\.\d+)?)', block, re.I):
            pt = (float(x), float(y))
            if not pairs or pairs[-1] != pt:
                pairs.append(pt)
        return pairs

    lines_found = []
    for block in blocks:
        pairs  = extract_coord_pairs(block)
        s_name = extract_name(block)
        style  = extract_style(block) or 'solid'

        if len(pairs) >= 2:
            x1, y1 = pairs[0]
            x2, y2 = pairs[1]
            x1, y1, x2, y2 = (int(v) if v == int(v) else v for v in (x1, y1, x2, y2))
            sl, b  = pts_to_slope(x1, y1, x2, y2)
            lines_found.append({
                'equation'      : s_name if s_name else 'line',
                'slope'         : sl,
                'intercept'     : b,
                'points'        : [(x1, y1), (x2, y2)],
                'linestyle'     : linestyles_map.get(style, '-'),
                'style_explicit': extract_style(block) is not None,
            })
        else:
            eq_m = re.search(r'y\s*=\s*(-?\s*\d*\.?\d*\s*x\s*(?:[+\-]\s*\d+\.?\d*)?|-?\d+\.?\d*)', block, re.I)
            if eq_m:
                result = parse_equation(eq_m.group(0))
                if result:
                    sl, b = result
                    x_lo, x_hi = params['x_min'], params['x_max']
                    pt1 = (int(x_lo), int(sl*x_lo+b) if (sl*x_lo+b) == int(sl*x_lo+b) else sl*x_lo+b)
                    pt2 = (int(x_hi), int(sl*x_hi+b) if (sl*x_hi+b) == int(sl*x_hi+b) else sl*x_hi+b)
                    lines_found.append({
                        'equation'      : s_name if s_name else eq_m.group(0).strip(),
                        'slope'         : sl,
                        'intercept'     : b,
                        'points'        : [pt1, pt2],
                        'linestyle'     : linestyles_map.get(style, '-'),
                        'style_explicit': extract_style(block) is not None,
                    })

    ds = get_dataset_styles(len(lines_found), params['color_mode'])
    for idx, line in enumerate(lines_found):
        line['color'] = ds[idx][0]
        if params['color_mode'] != 'color' and not line.get('style_explicit'):
            line['linestyle'] = ds[idx][1]
    params['lines'] = lines_found

    int_m = re.search(
        r'intersect(?:ion)?(?:\s+point)?\s*:?\s*(?:at|near)?\s*(?:approximately\s+)?(?:the\s+point\s+)?\(?\s*(-?\d+(?:\.\d+)?)\s*,\s*(-?\d+(?:\.\d+)?)\)?',
        p_flat, re.I)
    if int_m:
        x, y = float(int_m.group(1)), float(int_m.group(2))
        params['intersection'] = (int(x) if x == int(x) else x, int(y) if y == int(y) else y)
    return params


def _synthesise_scatter_points(description, x_min, x_max, y_min, y_max, seed=42):
    import random
    rng = random.Random(seed)
    n   = 10
    xs  = [x_min + (x_max - x_min) * i / (n - 1) for i in range(n)]
    d   = description.lower()

    def jitter(v, lo, hi, amount=0.7):
        return max(lo, min(hi, v + rng.uniform(-amount, amount) * (hi - lo) / n))

    if re.search(r'no\s+clear\s+pattern|random|no\s+association|scattered\s+with\s+no', d):
        pts = [(xs[i], y_min + rng.random() * (y_max - y_min)) for i in range(n)]
    elif re.search(r'downward|negative|upper\s+left\s+to\s+lower\s+right|falls?|decreas', d):
        pts = [(xs[i], jitter(y_max - (y_max - y_min) * i / (n - 1), y_min, y_max)) for i in range(n)]
    elif re.search(r'upward|positive|lower\s+left\s+to\s+upper\s+right|ris(es?|ing)|increas|nearly\s+straight', d):
        pts = [(xs[i], jitter(y_min + (y_max - y_min) * i / (n - 1), y_min, y_max, 0.3)) for i in range(n)]
    elif re.search(r'revers', d):
        pts = [(jitter(y_min + (y_max - y_min) * i / (n - 1), x_min, x_max, 0.3), xs[i]) for i in range(n)]
    else:
        pts = [(xs[i], jitter(y_min + (y_max - y_min) * i / (n - 1), y_min, y_max)) for i in range(n)]

    return [(round(x, 1), round(y, 1)) for x, y in pts]


def _extract_axis_labels_from_intro(text):
    m = re.search(
        r'scatter\s+plots?\s+(?:showing|of|for|comparing)\s+'
        r'([a-z][a-z\s]+?)\s+and\s+([a-z][a-z\s]+?)[\.,]',
        text, re.I
    )
    if m:
        return to_sentence_case(m.group(1).strip()), to_sentence_case(m.group(2).strip())
    return None, None


def parse_scatter_prompt(prompt):
    p = prompt.replace('\n', ' ').strip()
    grade_band = extract_grade_band(p)
    params = {
        'x_label': 'x', 'x_min': 0, 'x_max': 10,
        'y_label': 'y', 'y_min': 0, 'y_max': 10,
        'plots': {}, 'grade_band': grade_band,
        'color_mode': resolve_color_mode(grade_band),
        'plot_descriptions': {},
    }

    xl = re.search(r'(?:x-axis|horizontal\s+axis)\s*(?:labeled?|as)?\s*"([^"]+)"', p, re.I)
    yl = re.search(r'(?:y-axis|vertical\s+axis)\s*(?:labeled?|as)?\s*"([^"]+)"', p, re.I)
    if xl: params['x_label'] = to_sentence_case(xl.group(1))
    if yl: params['y_label'] = to_sentence_case(yl.group(1))

    if not xl or not yl:
        ix, iy = _extract_axis_labels_from_intro(p)
        if ix and not xl: params['x_label'] = ix
        if iy and not yl: params['y_label'] = iy

    xr = re.search(r'(?:x-axis|horizontal\s+axis)[^.]*?from\s+(-?\d+(?:\.\d+)?)\s+to\s+(-?\d+(?:\.\d+)?)', p, re.I)
    yr = re.search(r'(?:y-axis|vertical\s+axis)[^.]*?from\s+(-?\d+(?:\.\d+)?)\s+to\s+(-?\d+(?:\.\d+)?)', p, re.I)
    if xr: params['x_min'], params['x_max'] = float(xr.group(1)), float(xr.group(2))
    if yr: params['y_min'], params['y_max'] = float(yr.group(1)), float(yr.group(2))

    plot_starts = [
        (m.group(1).upper(), m.start())
        for m in re.finditer(r'\bPlot\s+([A-Za-z0-9])\b', p, re.I)
    ]
    plots = {}
    descriptions = {}

    for i, (label, start) in enumerate(plot_starts):
        end   = plot_starts[i + 1][1] if i + 1 < len(plot_starts) else len(p)
        block = p[start:end]

        paren_pts = re.findall(r'\(\s*(-?\d+(?:\.\d+)?)\s*,\s*(-?\d+(?:\.\d+)?)\s*\)', block)
        if paren_pts:
            pts = [(float(x), float(y)) for x, y in paren_pts]
            plots[label] = [(int(x) if x == int(x) else x, int(y) if y == int(y) else y) for x, y in pts]
            continue

        m_pts = re.search(r'points?\s+([\d.,\s\-]+)', block, re.I)
        if m_pts:
            nums = [float(n) for n in re.findall(r'-?\d+(?:\.\d+)?', m_pts.group(1))]
            pts  = list(zip(nums[0::2], nums[1::2]))
            if pts:
                plots[label] = [(int(x) if x == int(x) else x, int(y) if y == int(y) else y) for x, y in pts]
                continue

        desc = re.sub(r'^Plot\s+[A-Za-z0-9]\s+shows?\s+points?\s*[—\-]*\s*', '', block, flags=re.I).strip()
        descriptions[label] = desc
        pts = _synthesise_scatter_points(
            desc, params['x_min'], params['x_max'], params['y_min'], params['y_max'], seed=ord(label)
        )
        plots[label] = pts

    params['plots']             = plots
    params['plot_descriptions'] = descriptions
    return params


def parse_prompt(prompt, use_ai=False, ai_api_key=None):
    """
    Default: free regex parser.
    If use_ai is True and a key is supplied, try AI parsing first;
    on any failure, silently fall back to the regex parser.
    Returns (params, img_type, parse_method).
    """
    if use_ai and ai_api_key:
        data, err = ai_parse_prompt(prompt, ai_api_key)
        if data:
            img_type = data.get('type')
            if img_type == 'line_graph':
                return normalize_ai_line_params(data), img_type, 'AI'
            elif img_type == 'scatter_grid':
                return normalize_ai_scatter_params(data), img_type, 'AI'
        # fall through to regex on any failure

    img_type = detect_prompt_type(prompt)
    if img_type == 'line_graph':
        return parse_line_prompt(prompt), img_type, 'rule-based (free)'
    elif img_type == 'scatter_grid':
        return parse_scatter_prompt(prompt), img_type, 'rule-based (free)'
    return None, img_type, 'rule-based (free)'


# ============================================================
# IMAGE RENDERING
# ============================================================
def generate_line_image(params):
    x_min, x_max = params['x_min'], params['x_max']
    y_min, y_max = params['y_min'], params['y_max']
    pad = STYLE_RULES['graphs']['axis_pad_units']
    all_points = [pt for line in params['lines'] for pt in line.get('points', [])]
    if params.get('intersection'):
        all_points.append(params['intersection'])
    if all_points:
        x_min = min(x_min, min(p[0] for p in all_points))
        x_max = max(x_max, max(p[0] for p in all_points))
        y_min = min(y_min, min(p[1] for p in all_points))
        y_max = max(y_max, max(p[1] for p in all_points))
    plt.rcParams['font.family'] = resolve_font(params.get('grade_band'))
    xp_min, xp_max = x_min - pad, x_max + pad
    yp_min, yp_max = y_min - pad, y_max + pad
    fig, ax = plt.subplots(figsize=(6, 6), facecolor='white')
    ax.set_facecolor('white')
    ax.grid(True, linestyle='--',
            linewidth=STYLE_RULES['graphs']['gridline_width_pt'],
            color=STYLE_RULES['graphs']['gridline_color'], alpha=0.6, zorder=0)
    ax.axhline(0, color='black', linewidth=0.8, zorder=1)
    ax.axvline(0, color='black', linewidth=0.8, zorder=1)
    for spine in ax.spines.values():
        spine.set_visible(False)
    for xy, xyt in [((xp_max,0),(xp_max-0.01,0)), ((xp_min,0),(xp_min+0.01,0)),
                    ((0,yp_max),(0,yp_max-0.01)), ((0,yp_min),(0,yp_min+0.01))]:
        ax.annotate('', xy=xy, xytext=xyt, arrowprops=dict(arrowstyle='->', color='black', lw=0.8))

    x_full = np.linspace(xp_min, xp_max, 1000)

    def crosshair(xi, yi, color):
        ax.plot([xi-0.15, xi+0.15], [yi, yi], color=color, linewidth=1.0, zorder=5)
        ax.plot([xi, xi], [yi-0.15, yi+0.15], color=color, linewidth=1.0, zorder=5)

    offsets = [(8,6),(8,-18),(-55,6),(-55,-18)]
    near_left = xp_min + (xp_max - xp_min) * 0.12
    for idx, line in enumerate(params['lines']):
        if line['slope'] is None: continue
        y_vals = line['slope'] * x_full + line['intercept']
        mask   = (y_vals >= yp_min) & (y_vals <= yp_max)
        ax.plot(x_full[mask], y_vals[mask], color=line['color'],
                linewidth=STYLE_RULES['graphs']['line_width_pt'],
                linestyle=line['linestyle'], label=line['equation'], zorder=2)
        for pidx, (xi, yi) in enumerate(line['points']):
            if params.get('intersection') and (xi, yi) == params['intersection']: continue
            crosshair(xi, yi, line['color'])
            offset = (10, 6) if (xi <= near_left and pidx % 2 == 0) else \
                     (10, -16) if xi <= near_left else offsets[(idx*2+pidx) % len(offsets)]
            ax.annotate(f'({xi},{yi})', xy=(xi, yi), xytext=offset,
                        textcoords='offset points', fontsize=9,
                        fontfamily='STIXGeneral', color=line['color'],
                        bbox=dict(boxstyle='round,pad=0.2', fc='white', ec='none', alpha=0.95), zorder=6)
    if params['intersection']:
        ix, iy = params['intersection']
        crosshair(ix, iy, 'black')
        ax.annotate(f'({ix},{iy})', xy=(ix, iy), xytext=(8, 8),
                    textcoords='offset points', fontsize=9, fontfamily='STIXGeneral', color='black',
                    bbox=dict(boxstyle='round,pad=0.2', fc='white', ec='none', alpha=0.95), zorder=7)
    x_step = max(1, round((x_max - x_min) / 8))
    y_step = max(1, round((y_max - y_min) / 8))
    ax.set_xticks(np.arange(x_min, x_max + x_step, x_step))
    ax.set_yticks(np.arange(y_min, y_max + y_step, y_step))
    ax.tick_params(labelsize=9)
    ax.set_xlabel(params['x_label'], fontsize=11, style='italic', labelpad=6)
    ax.set_ylabel(params['y_label'], fontsize=11, style='italic', labelpad=6)
    if params.get('title'):
        ax.set_title(params['title'], fontsize=14, fontweight='bold', pad=12)
    if len(params['lines']) > 1:
        ax.legend(loc='upper right', fontsize=9, frameon=True, framealpha=0.9, edgecolor='none')
    plt.tight_layout(rect=[0, 0.02, 1, 1])
    ax.set_xlim(xp_min, xp_max)
    ax.set_ylim(yp_min, yp_max)
    buf = io.BytesIO()
    plt.savefig(buf, dpi=300, bbox_inches='tight', facecolor='white', format='png')
    plt.close(fig)
    buf.seek(0)
    return buf


def generate_scatter_grid(params):
    plots      = params['plots']
    labels     = sorted(plots.keys())
    color_mode = params.get('color_mode', 'grayscale')
    ds         = get_dataset_styles(len(labels), color_mode)
    n_rows = 2 if len(labels) > 2 else 1
    n_cols = 2 if len(labels) > 1 else 1
    fig, axes = plt.subplots(n_rows, n_cols, figsize=(5*n_cols, 5*n_rows), facecolor='white', squeeze=False)
    fig.subplots_adjust(hspace=0.45, wspace=0.35)
    axes_flat = axes.flatten()
    for ax, label, style in zip(axes_flat, labels, ds):
        points = plots[label]
        xs = [p[0] for p in points]
        ys = [p[1] for p in points]
        color = style[0] if color_mode == 'color' else 'black'
        ax.scatter(xs, ys, color=color, marker='o', s=60,
                   edgecolors='black' if color_mode != 'color' else color,
                   linewidths=0.8, zorder=3)
        ax.grid(True, linestyle='--', linewidth=0.5, color='#cccccc', alpha=0.7)
        px_max = max(params['x_max'], max(xs)+2) if xs else params['x_max']
        py_max = max(params['y_max'], max(ys)+2) if ys else params['y_max']
        ax.set_xlim(params['x_min'], px_max)
        ax.set_ylim(params['y_min'], py_max)
        ax.set_xticks(np.arange(params['x_min'], px_max+1, max(1, (px_max - params['x_min']) // 7)))
        ax.set_yticks(np.arange(params['y_min'], py_max+1, max(1, (py_max - params['y_min']) // 6)))
        ax.tick_params(labelsize=8)
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.set_xlabel(params['x_label'], fontsize=9, labelpad=4)
        ax.set_ylabel(params['y_label'], fontsize=9, labelpad=4)
        ax.set_title(f'Plot {label}', fontsize=14, fontweight='bold', loc='left', pad=8)
    for i in range(len(labels), len(axes_flat)):
        axes_flat[i].set_visible(False)
    buf = io.BytesIO()
    plt.savefig(buf, dpi=300, bbox_inches='tight', facecolor='white', format='png')
    plt.close(fig)
    buf.seek(0)
    return buf


def render_image(params, img_type):
    if img_type == 'line_graph':    return generate_line_image(params)
    if img_type == 'scatter_grid':  return generate_scatter_grid(params)
    raise ValueError(f'Unknown type: {img_type}')


# ============================================================
# DOCX HELPERS
# ============================================================
def extract_full_text(para):
    parts = []
    for elem in para._element.iter():
        tag = elem.tag.split('}')[-1] if '}' in elem.tag else elem.tag
        if tag == 't' and elem.text:
            parts.append(elem.text)
        elif tag in ('begChr', 'endChr'):
            val = elem.get('{http://schemas.openxmlformats.org/officeDocument/2006/math}val', '')
            if val: parts.append(val)
    return ''.join(parts)

def find_image_placeholders(doc):
    found = []
    for i, para in enumerate(doc.paragraphs):
        text = extract_full_text(para)
        if re.search(r'image\s*\(if any\)\s*:', text, re.IGNORECASE):
            has_content = len(text.strip()) > len('Image (if any):') + 2
            found.append({'para_index': i, 'text': text.strip(), 'has_content': has_content})
    return found

def replace_placeholder(doc, para_index, img_bytes):
    """Replace the target paragraph's XML with a fresh <w:p> containing only
    the image run — avoids leftover runs / broken element trees."""
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from copy        import deepcopy

    para   = doc.paragraphs[para_index]
    p_elem = para._element

    pPr = p_elem.find(qn('w:pPr'))
    pPr_copy = deepcopy(pPr) if pPr is not None else None

    new_p = OxmlElement('w:p')
    if pPr_copy is not None:
        new_p.append(pPr_copy)

    new_r = OxmlElement('w:r')
    new_p.append(new_r)

    tmp_para = doc.add_paragraph()
    tmp_run  = tmp_para.add_run()
    img_bytes.seek(0)
    tmp_run.add_picture(img_bytes, width=Inches(5.5))

    drawing = tmp_para._element.find('.//' + qn('w:drawing'))
    if drawing is not None:
        new_r.append(deepcopy(drawing))

    tmp_para._element.getparent().remove(tmp_para._element)
    p_elem.getparent().replace(p_elem, new_p)


def process_doc_bytes(doc_bytes, use_ai=False, ai_api_key=None):
    """Process a .docx bytes blob. Returns (updated_doc_bytes, results_list)."""
    doc          = Document(io.BytesIO(doc_bytes))
    placeholders = find_image_placeholders(doc)
    results      = []
    for ph in placeholders:
        raw    = ph['text']
        prompt = re.sub(r'^image\s*\(if any\)\s*:\s*', '', raw, flags=re.I).strip()
        entry  = {'placeholder': raw, 'para_index': ph['para_index'], 'prompt': prompt,
                  'img_type': None, 'parse_method': None, 'img_buf': None, 'error': None}
        if not prompt:
            entry['error'] = 'Empty prompt'
            results.append(entry)
            continue
        try:
            params, img_type, parse_method = parse_prompt(prompt, use_ai=use_ai, ai_api_key=ai_api_key)
            entry['img_type']     = img_type
            entry['parse_method'] = parse_method
            if img_type not in ('line_graph', 'scatter_grid'):
                entry['error'] = 'Could not detect graph type'
                results.append(entry)
                continue
            entry['img_buf'] = render_image(params, img_type)
        except Exception as e:
            entry['error'] = str(e)
        results.append(entry)

    # Embed in reverse paragraph order so earlier replacements don't shift later indices
    for r in sorted(results, key=lambda x: x['para_index'], reverse=True):
        if r['img_buf'] and not r['error']:
            r['img_buf'].seek(0)
            replace_placeholder(doc, r['para_index'], r['img_buf'])

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out.read(), results


# ============================================================
# GOOGLE DRIVE HELPERS
# ============================================================
def extract_folder_id(url_or_id):
    m = re.search(r'/folders/([a-zA-Z0-9_-]+)', url_or_id)
    if m: return m.group(1)
    if re.match(r'^[a-zA-Z0-9_-]{20,}$', url_or_id.strip()):
        return url_or_id.strip()
    return None

def get_drive_service(creds_json_str):
    import json as _json
    from google.oauth2 import service_account
    from googleapiclient.discovery import build
    info  = _json.loads(creds_json_str)
    creds = service_account.Credentials.from_service_account_info(
        info, scopes=['https://www.googleapis.com/auth/drive'])
    return build('drive', 'v3', credentials=creds)

def list_docx_in_folder(service, folder_id):
    q = (f"'{folder_id}' in parents and "
         "mimeType='application/vnd.openxmlformats-officedocument.wordprocessingml.document' "
         "and trashed=false")
    return service.files().list(q=q, fields='files(id,name)').execute().get('files', [])

def download_drive_file(service, file_id):
    from googleapiclient.http import MediaIoBaseDownload
    buf = io.BytesIO()
    dl  = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
    done = False
    while not done: _, done = dl.next_chunk()
    buf.seek(0)
    return buf.read()

def update_drive_file(service, file_id, data, mime):
    """Update an EXISTING Drive file in-place — uses the file owner's quota,
    not the service account's (which has none), so this avoids the
    storageQuotaExceeded error on CREATE."""
    from googleapiclient.http import MediaIoBaseUpload
    media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime)
    f = service.files().update(fileId=file_id, media_body=media, fields='id,webViewLink').execute()
    return f.get('webViewLink', '')


# ============================================================
# SIDEBAR
# ============================================================
with st.sidebar:
    st.title('⚙️ Settings')
    st.markdown('---')
    mode = st.radio('**Mode**', ['📄 Single Document', '📁 Batch (Google Drive)'])

    st.markdown('---')
    st.markdown('**Prompt Parsing**')
    use_ai = st.checkbox(
        '🤖 Use AI parsing (optional, costs money)',
        value=False,
        help='Off by default. The free rule-based parser handles most prompts. '
             'Only enable this if you have your own Anthropic API key and want '
             'extra-flexible prompt understanding.'
    )
    ai_api_key = None
    if use_ai:
        default_key = ''
        try: default_key = st.secrets.get('ANTHROPIC_API_KEY', '')
        except Exception: pass
        ai_api_key = st.text_input(
            'Anthropic API Key', value=default_key, type='password',
            help='Your own key from console.anthropic.com — billed to you, pay-as-you-go.'
        )
        st.caption('⚠️ This uses your API credits. Falls back to the free parser automatically if it fails.')

    if 'Batch' in mode:
        st.markdown('---')
        st.markdown('**Google Drive — Service Account JSON**')
        default_creds = ''
        try: default_creds = st.secrets.get('GOOGLE_CREDS_JSON', '')
        except Exception: pass
        creds_input = st.text_area(
            'Paste Service Account JSON',
            value=default_creds, height=140,
            placeholder='{"type": "service_account", ...}',
            help='Share your Drive folder with the service account email (Editor access).'
        )
        st.caption('🔒 Never stored or logged.')

    st.markdown('---')
    st.caption('Graph generation always runs free on the server. AI parsing is optional and off by default.')


# ============================================================
# SINGLE DOCUMENT MODE
# ============================================================
if '📄 Single' in mode:
    st.title('📐 Math Image Generator')

    with st.expander('📖 How to use this app (click to open)', expanded=False):
        st.markdown("""
**Step-by-step guide:**

1. **Prepare your Word document (.docx)**
   - Wherever you want a graph to appear, type this on its own line:
     ```
     Image (if any): <your graph description here>
     ```
   - Example: `Image (if any): the line y = 2x + 1, passing through (0,1) and (2,5)`

2. **Upload your document** using the "Upload .docx" button below.
3. **Select the placeholder** — the app detects all `Image (if any):` lines.
4. **Paste your graph description** into the "Image prompt" box.
5. **Click "Generate Image"** — a preview appears on the right.
6. **Download** the image (PNG) or the full updated document.

---
**What kinds of graphs work?**

| Graph type | Example prompt |
|---|---|
| Line graph | `the line y = 2x + 1, passing through (0,1) and (2,5)` |
| Multiple lines | `Line A: (0,1) and (2,5) (solid). Line B: (0,3) and (4,3) (dashed).` |
| Scatter plots | `Plot A: (1,2), (3,4), (5,6). Plot B: (2,5), (4,8).` |

**Color vs grayscale?** Add `grade 3-5` or `elementary` for color; `grade 6-8`/`grade 9-12` for grayscale.

**AI parsing (optional, sidebar)** — off by default, free regex parser is used unless you turn it on and supply your own API key.
        """)

    st.markdown('---')
    col1, col2 = st.columns([1, 1])

    with col1:
        st.subheader('1 — Upload Document')
        uploaded = st.file_uploader('Upload .docx', type=['docx'])
        placeholders = []

        if uploaded:
            doc_obj      = Document(uploaded)
            placeholders = find_image_placeholders(doc_obj)
            if placeholders:
                st.success(f'✅ Found {len(placeholders)} placeholder(s)')
            else:
                st.warning('No `Image (if any):` placeholders found in this document.')

        st.subheader('2 — Select Placeholder')
        selected_idx = None
        if placeholders:
            selected_idx = st.selectbox(
                'Choose a placeholder to fill',
                options=range(len(placeholders)),
                format_func=lambda i: f"[Para {placeholders[i]['para_index']}] {placeholders[i]['text'][:70]}..."
            )
        else:
            st.info('Upload a document to see placeholders.')

        if placeholders and selected_idx is not None:
            debug_raw    = placeholders[selected_idx]['text']
            debug_prompt = re.sub(r'^image\s*\(if any\)\s*:\s*', '', debug_raw, flags=re.I).strip()
            with st.expander('🔍 Debug: prompt extracted from document', expanded=True):
                st.code(debug_prompt or '(empty — nothing found after "Image (if any):")')

        st.subheader('3 — Paste Prompt')
        prompt = st.text_area(
            'Graph description',
            height=180,
            placeholder='Example: the line y = 2x + 1, passing through (0,1) and (2,5), x-axis from 0 to 6, y-axis from 0 to 12'
        )

    with col2:
        st.subheader('Preview & Download')
        if st.button('▶️ Generate Image', type='primary', use_container_width=True):
            if not uploaded:
                st.error('Please upload a .docx file first.')
            elif not prompt.strip():
                st.error('Please paste a graph description in the box on the left.')
            elif selected_idx is None:
                st.error('No placeholder selected.')
            else:
                with st.spinner('Generating…'):
                    try:
                        params, img_type, parse_method = parse_prompt(
                            prompt.strip(), use_ai=use_ai, ai_api_key=ai_api_key)
                        st.info(f'Detected type: **{img_type}** · Parsed with: **{parse_method}**')

                        if img_type == 'line_graph' and (not params or not params.get('lines')):
                            st.error(
                                '⚠️ No lines found in your prompt.\n\n'
                                'Make sure it includes either:\n'
                                '- An equation like `y = 2x + 1`, or\n'
                                '- Two coordinate points like `(0,1)` and `(2,5)`'
                            )
                        elif img_type == 'scatter_grid' and (not params or not params.get('plots')):
                            st.error(
                                '⚠️ No scatter data found.\n\n'
                                'Your prompt should include sections like:\n'
                                '`Plot A: (1,2), (3,4), (5,6)`'
                            )
                        elif img_type not in ('line_graph', 'scatter_grid'):
                            st.error(
                                '⚠️ Could not figure out what type of graph to draw.\n\n'
                                'Try including words like "line", "y =", "scatter", or "Plot A" in your description.'
                            )
                        else:
                            img_buf = render_image(params, img_type)
                            st.image(img_buf, caption='Generated graph', use_container_width=True)

                            img_buf.seek(0)
                            uploaded.seek(0)
                            doc2 = Document(uploaded)
                            img_buf.seek(0)
                            replace_placeholder(doc2, placeholders[selected_idx]['para_index'], img_buf)
                            doc_out = io.BytesIO()
                            doc2.save(doc_out)
                            doc_out.seek(0)
                            img_buf.seek(0)

                            st.download_button('⬇️ Download Image (.png)', data=img_buf,
                                               file_name='math_image.png', mime='image/png',
                                               use_container_width=True)
                            st.download_button('⬇️ Download Updated Document (.docx)', data=doc_out,
                                               file_name='output_with_image.docx',
                                               mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                                               use_container_width=True)
                            st.success('✅ Done! Download your files above.')
                    except Exception as e:
                        st.error(f'Something went wrong: {e}')
                        import traceback; st.code(traceback.format_exc())


# ============================================================
# BATCH (GOOGLE DRIVE) MODE
# ============================================================
else:
    st.title('📁 Batch Mode — Google Drive')

    with st.expander('📖 How to use Batch Mode (click to open)', expanded=False):
        st.markdown("""
**What this does:** Automatically processes an entire folder of Word documents from Google Drive,
embeds graphs into each one, and lets you download all updated files in a single ZIP file.

**Setup (one-time):**
1. Go to [Google Cloud Console](https://console.cloud.google.com) and enable the **Google Drive API**.
2. Create a **Service Account** and download the JSON key file.
3. **Share your Drive folder** with the service account email (give it **Editor** access).
4. Paste the JSON key contents into the sidebar on the left.

**Then:**
1. Paste your Google Drive folder link below.
2. Click **Process All Documents**.
3. Wait for processing to finish, then download the ZIP file.

The ZIP will contain:
- All updated `.docx` files (with graphs embedded)
- All graphs as separate `.png` files
        """)

    st.markdown('---')
    folder_url = st.text_input(
        '🔗 Google Drive Folder Link',
        placeholder='https://drive.google.com/drive/folders/XXXXXXXXXX'
    )

    show_batch_debug = st.checkbox(
        '🔍 Show debug info (extracted prompts per document)',
        value=False,
        help='When checked, after processing you can expand each document to see exactly '
             'what prompt text was extracted from every placeholder, what type was detected, '
             'and how it was parsed.'
    )

    if st.button('🚀 Process All Documents', type='primary'):
        creds_str = creds_input.strip() if 'creds_input' in dir() else ''
        if not creds_str:
            st.error('Please paste your Service Account JSON in the sidebar on the left first.')
            st.stop()
        if not folder_url.strip():
            st.error('Please enter a Google Drive folder link.')
            st.stop()

        folder_id = extract_folder_id(folder_url.strip())
        if not folder_id:
            st.error(
                'Could not read the folder ID from that link.\n\n'
                'The link should look like: `https://drive.google.com/drive/folders/XXXXX`'
            )
            st.stop()

        with st.spinner('Connecting to Google Drive…'):
            try:
                service = get_drive_service(creds_str)
                files   = list_docx_in_folder(service, folder_id)
            except Exception as e:
                st.error(f'Could not connect to Google Drive: {e}')
                st.stop()

        if not files:
            st.warning('No .docx files found in that folder.')
            st.stop()

        st.success(f'Found **{len(files)}** document(s). Processing…')

        upload_back = st.checkbox(
            '☁️ Also save updated files back to Google Drive '
            '(replaces the original .docx in-place — no storage quota needed)',
            value=True
        )

        progress     = st.progress(0)
        status       = st.empty()
        summary      = []
        debug_blocks = []   # (fname, list_of_result_dicts) — used if show_batch_debug
        zip_buf      = io.BytesIO()

        DOCX_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for idx, f in enumerate(files):
                fname   = f['name']
                file_id = f['id']
                stem    = Path(fname).stem
                status.markdown(f'⏳ Working on **{fname}** ({idx+1} of {len(files)})…')
                try:
                    doc_bytes = download_drive_file(service, file_id)
                    updated_bytes, results = process_doc_bytes(
                        doc_bytes, use_ai=use_ai, ai_api_key=ai_api_key)

                    if not results:
                        summary.append(f'⚠️ **{fname}** — no placeholders found, skipped')
                        progress.progress((idx+1)/len(files))
                        continue

                    if show_batch_debug:
                        debug_blocks.append((fname, results))

                    zf.writestr(f'Updated_Docs/{fname}', updated_bytes)
                    for r_idx, r in enumerate(results):
                        if r['img_buf'] and not r['error']:
                            r['img_buf'].seek(0)
                            zf.writestr(f'Images/{stem}_image_{r_idx+1}.png', r['img_buf'].read())

                    drive_saved = False
                    if upload_back:
                        try:
                            update_drive_file(service, file_id, updated_bytes, DOCX_MIME)
                            drive_saved = True
                        except Exception as ue:
                            summary.append(f'  ⚠️ Drive update failed for **{fname}**: `{ue}`')

                    ok  = sum(1 for r in results if r['img_buf'] and not r['error'])
                    err = sum(1 for r in results if r['error'])
                    drive_note = ' · ☁️ updated in Drive' if drive_saved else ''
                    summary.append(
                        f'✅ **{fname}** — {ok} graph(s) embedded{drive_note}'
                        + (f', {err} skipped (prompt not recognised)' if err else '')
                    )
                except Exception as e:
                    summary.append(f'❌ **{fname}** — error: {e}')
                progress.progress((idx+1)/len(files))

        status.empty()
        zip_buf.seek(0)

        st.markdown('---')
        st.subheader('✅ All done!')
        for line in summary:
            st.markdown(line)

        if upload_back:
            st.info(
                '☁️ **Drive update:** The original `.docx` files in your folder have been updated '
                'in-place with graphs embedded. Open them directly in Google Drive to check.'
            )

        # ---- Batch debug view ----
        if show_batch_debug and debug_blocks:
            st.markdown('---')
            st.subheader('🔍 Debug: prompts extracted per document')
            for fname, results in debug_blocks:
                with st.expander(f'📄 {fname}  —  {len(results)} placeholder(s)', expanded=False):
                    for i, r in enumerate(results):
                        st.markdown(f"**Placeholder {i+1}** (paragraph {r['para_index']})")
                        st.code(r.get('prompt') or '(empty)')
                        status_line = f"Type: `{r['img_type']}`"
                        if r.get('parse_method'):
                            status_line += f" · Parsed with: `{r['parse_method']}`"
                        if r['error']:
                            status_line += f" · ❌ Error: {r['error']}"
                        else:
                            status_line += " · ✅ Image generated"
                        st.caption(status_line)
                        st.markdown('---')

        st.download_button(
            '⬇️ Download All Updated Files (ZIP)',
            data=zip_buf,
            file_name='Updated_Docs_with_Graphs.zip',
            mime='application/zip',
            use_container_width=True
        )
